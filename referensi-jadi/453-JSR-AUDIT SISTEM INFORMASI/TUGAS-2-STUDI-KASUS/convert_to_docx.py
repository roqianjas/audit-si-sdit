"""
Konversi Studi Kasus COBIT (Tugas 2) dari Markdown ke Word (.docx)
Format Akademis: Margin 4-3-3-3, TNR 12pt, Spasi 1.5
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import os

BASE_DIR = r"c:\Users\adminlocal\Documents\UNM\UNM-Semester-8\453-JSR-AUDIT SISTEM INFORMASI\TUGAS-2-STUDI-KASUS"
OUTPUT_PATH = os.path.join(BASE_DIR, "Tugas2-Studi_Kasus_COBIT-Roki_Anjas.docx")

# === SETUP ===
doc = Document()
for section in doc.sections:
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(13)
doc.styles['Heading 3'].font.size = Pt(12)


# === HELPERS ===
def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def normal(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
        else:
            r = p.add_run(part)
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold
    return p

def list_item(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(strip_md(text))
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(strip_md(str(val)))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def parse_md_table(text):
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    lines = [l for l in lines if not re.match(r'^\|[-\s|:]+\|$', l)]
    if not lines: return [], []
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    rows = [[c.strip() for c in l.split('|')[1:-1]] for l in lines[1:]]
    return headers, [r for r in rows if r]

def process_md_file(filepath, skip_h1=True):
    content = open(filepath, 'r', encoding='utf-8').read()
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line.strip() == '---':
            i += 1; continue
        if skip_h1 and line.startswith('# '):
            i += 1; continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(strip_md(line[5:].strip()))
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue

        bold_match = re.match(r'^\*\*(.+?)\*\*$', line.strip())
        if bold_match and not line.strip().startswith('|'):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(bold_match.group(1))
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue

        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i]); i += 1
            headers, rows = parse_md_table('\n'.join(table_lines))
            if headers and rows:
                add_table(headers, rows)
            continue

        if line.strip().startswith('>'):
            i += 1; continue

        # Code block
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1.25)
                r = p.add_run(cl)
                r.font.name = 'Courier New'; r.font.size = Pt(10)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        if re.match(r'^\d+\.', line.strip()):
            list_item(line.strip()); i += 1; continue

        if line.strip().startswith('- ') or line.strip().startswith('• '):
            text = line.strip().lstrip('-•').strip()
            list_item(f"• {strip_md(text)}"); i += 1; continue

        text = line.strip()
        if text and not text.startswith(('#', '|', '!', '>')):
            normal(text)
        i += 1


# ============================================================
# COVER
# ============================================================
centered("TUGAS 2", 16, bold=True, space_after=0)
centered("STUDI KASUS IMPLEMENTASI COBIT", 16, bold=True, space_after=18)
centered("", 8)
centered("MATA KULIAH AUDIT SISTEM INFORMASI (453)", 14, bold=True, space_after=18)
centered("", 12)
centered("Disusun Oleh:", 12, bold=True, space_after=4)
centered("Roki Anjas — 11250066", 12, bold=False, space_after=18)
centered("", 8)
centered("Dosen Pengampu:", 12, bold=True, space_after=0)
centered("Juarni Siregar, S.Pd., M.Kom", 12, bold=True, space_after=18)
centered("", 8)
centered("PROGRAM STUDI SISTEM INFORMASI (S1)", 12, bold=True, space_after=0)
centered("FAKULTAS TEKNOLOGI INFORMASI", 12, bold=True, space_after=0)
centered("UNIVERSITAS NUSA MANDIRI", 12, bold=True, space_after=0)
p_last = centered("2026", 12, bold=True)
p_last.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)


# ============================================================
# STUDI KASUS 1 — COBIT 4
# ============================================================
centered("STUDI KASUS 1", 14, bold=True, space_after=0)
centered("IMPLEMENTASI COBIT 4", 14, bold=True, space_after=24)

process_md_file(os.path.join(BASE_DIR, "Studi_Kasus_COBIT4.md"))
doc.add_page_break()


# ============================================================
# STUDI KASUS 2 — COBIT 5
# ============================================================
centered("STUDI KASUS 2", 14, bold=True, space_after=0)
centered("IMPLEMENTASI COBIT 5", 14, bold=True, space_after=24)

process_md_file(os.path.join(BASE_DIR, "Studi_Kasus_COBIT5.md"))
doc.add_page_break()


# ============================================================
# STUDI KASUS 3 — COBIT 2019
# ============================================================
centered("STUDI KASUS 3", 14, bold=True, space_after=0)
centered("IMPLEMENTASI COBIT 2019", 14, bold=True, space_after=24)

process_md_file(os.path.join(BASE_DIR, "Studi_Kasus_COBIT2019.md"))


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"✅ Tugas 2 berhasil disimpan: {OUTPUT_PATH}")
print(f"📄 Format: TNR 12pt, Spasi 1.5, Margin 4-3-3-3 cm")
print(f"📝 Isi: 3 Studi Kasus (COBIT 4, COBIT 5, COBIT 2019)")
