"""
Konversi Laporan Project Audit SI SD IT Al-huda dari Markdown ke Word (.docx)
Format Akademis: Margin 4-3-2.5-2.5, TNR 12pt, Spasi 1.5
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re
import os

BASE_DIR = "/www/wwwroot/audit-si-sdit/TUGAS-AUDIT-SDIT"
LAPORAN_DIR = os.path.join(BASE_DIR, "02_Laporan")
CHARTS_DIR = os.path.join(BASE_DIR, "03_Data", "charts")
OUTPUT_PATH = os.path.join(LAPORAN_DIR, "Laporan-Audit-SI-SDIT_Al-huda.docx")
LOGO_PATH = os.path.join(BASE_DIR, "..", "assets", "logo-unm.png")

# === GAMBAR MAP: caption keyword -> (file, width_cm) ===
GAMBAR_MAP = {
    'Gambar 2.1': ('gambar_4_metodologi.png', 14),
    'Gambar 2.2': ('gambar_1_radar_capability.png', 13),
    'Gambar 2.3': ('gambar_2_gap_analysis.png', 14),
    'Gambar 2.4': ('gambar_3_ketercapaian.png', 14),
}

# === SETUP ===
doc = Document()

def setup_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.footer_distance = Cm(0.3)
    section.header_distance = Cm(1.25)

for section in doc.sections:
    setup_section(section)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 3'].font.size = Pt(12)


# === PAGINATION HELPERS ===

def add_page_field(paragraph, fmt="arabic"):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    if fmt == "roman":
        instrText.text = ' PAGE \\* roman '
    else:
        instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)
    for r in paragraph.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def set_page_number_start(section, start=1, fmt="decimal"):
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    pgNumType.set(qn('w:fmt'), fmt)
    sectPr.append(pgNumType)

def clear_header_footer(section):
    for attr in ['header', 'footer', 'first_page_header', 'first_page_footer']:
        try:
            obj = getattr(section, attr)
            obj.is_linked_to_previous = False
            for p in obj.paragraphs:
                p.text = ''
        except:
            pass

def setup_cover_section(section):
    setup_section(section)
    clear_header_footer(section)

def setup_front_matter_section(section, start_page=2):
    setup_section(section)
    set_page_number_start(section, start=start_page, fmt="lowerRoman")
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="roman")
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ''

def setup_bab_section(section, restart_at=None):
    setup_section(section)
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    if restart_at is not None:
        set_page_number_start(section, start=restart_at, fmt="decimal")
    else:
        for pgNum in sectPr.findall(qn('w:pgNumType')):
            sectPr.remove(pgNum)
    section.first_page_footer.is_linked_to_previous = False
    fp = section.first_page_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="arabic")
    section.first_page_header.is_linked_to_previous = False
    for p in section.first_page_header.paragraphs:
        p.text = ''
    section.header.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(hp, fmt="arabic")
    section.footer.is_linked_to_previous = False
    for p in section.footer.paragraphs:
        p.text = ''

def setup_daftar_pustaka_section(section):
    setup_section(section)
    section.different_first_page_header_footer = False
    sectPr = section._sectPr
    for pgNum in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(pgNum)
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="arabic")
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ''

def add_new_section(section_type=WD_SECTION.NEW_PAGE):
    doc.add_section(section_type)
    return doc.sections[-1]

def add_toc_entry(text, page_text, is_bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(1.25)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = is_bold
    r2 = p.add_run(f"\t{page_text}")
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)


# === HELPERS ===
def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def normal(text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    
    # Handle inline bold and italic
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.name = 'Times New Roman'; r.font.size = Pt(12)
            r.font.bold = True
        else:
            # Handle italic within non-bold parts
            italic_parts = re.split(r'(\*[^*]+?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    r = p.add_run(ipart[1:-1])
                    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
                    r.font.italic = True
                else:
                    r = p.add_run(ipart)
                    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
                    r.font.bold = bold; r.font.italic = italic
    return p

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def list_item(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    
    # Handle inline bold
    parts = re.split(r'(\*\*.*?\*\*)', strip_md(text))
    for part in parts:
        r = p.add_run(part)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # Full width (100%)
    tbl = table._tbl
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tbl.tblPr.append(tblW)
    # Narrow 'No' columns
    for c_idx, h in enumerate(headers):
        if h.strip().lower() in ["no", "no.", "level"]:
            table.columns[c_idx].width = Cm(1.5)
    # Header row
    hdr_row = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell_para.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
        cell_para.paragraph_format.line_spacing = 1.0
        cell_para.paragraph_format.space_after = Pt(2)
        cell_para.paragraph_format.space_before = Pt(2)
        if h.strip().lower() in ["no", "no.", "level"]:
            cell.width = Cm(1.5)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cell_para.add_run(strip_md(str(val)))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            cell_para.paragraph_format.line_spacing = 1.0
            cell_para.paragraph_format.space_after = Pt(2)
            cell_para.paragraph_format.space_before = Pt(2)
            if headers[ci].strip().lower() in ["no", "no.", "level"]:
                cell.width = Cm(1.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_image(image_path, caption, width_cm=14):
    """Insert an image with caption below it."""
    if not os.path.exists(image_path):
        print(f"[WARN] Gambar tidak ditemukan: {image_path}")
        return
    # Image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(image_path, width=Cm(width_cm))
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    r_cap = p_cap.add_run(caption)
    r_cap.font.name = 'Times New Roman'
    r_cap.font.size = Pt(10)
    r_cap.font.italic = True

def parse_md_table(text):
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    lines = [l for l in lines if not re.match(r'^\|[-\s|:]+\|$', l)]
    if not lines: return [], []
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    rows = [[c.strip() for c in l.split('|')[1:-1]] for l in lines[1:]]
    return headers, [r for r in rows if r]

def process_md_content(content, skip_title=True):
    """Process markdown content and add to document."""
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line.strip() == '---':
            i += 1; continue
        # Skip table separator rows (e.g. |---|:---:|)
        if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
            i += 1; continue
        # Skip stray separator fragments (e.g. -:|, --:|)
        if re.match(r'^[-:|\s]+$', line.strip()) and len(line.strip()) <= 10:
            i += 1; continue
        if skip_title and (line.startswith('# BAB') or line.startswith('# DAFTAR')):
            i += 1; continue

        # Headings
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(strip_md(title))
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue

        # Bold subtitle on its own line
        bold_match = re.match(r'^\*\*(.+?)\*\*$', line.strip())
        if bold_match and not line.strip().startswith('|'):
            text = bold_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            if text.startswith('Tabel') or text.startswith('Gambar'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            else:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue

        # Table
        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i]); i += 1
            headers, rows = parse_md_table('\n'.join(table_lines))
            if headers and rows:
                add_table(headers, rows)
            continue

        # Blockquote — check for image reference
        if line.strip().startswith('>'):
            bq_text = line.strip().lstrip('>').strip()
            # Check if this is a Gambar reference: > **Gambar X.X** Caption
            gambar_match = re.match(r'^\*\*(.+?)\*\*\s*(.*)', bq_text)
            if gambar_match:
                gambar_key = gambar_match.group(1).strip()
                gambar_caption = gambar_match.group(2).strip()
                full_caption = f"{gambar_key} \u2014 {gambar_caption}" if gambar_caption else gambar_key
                # Look up in GAMBAR_MAP
                for key, (filename, width) in GAMBAR_MAP.items():
                    if key in gambar_key:
                        add_image(
                            os.path.join(CHARTS_DIR, filename),
                            full_caption,
                            width_cm=width
                        )
                        break
            i += 1; continue

        # Sub-item with letter (a., b., etc.)
        if re.match(r'^[a-z]\.\s', line.strip()):
            list_item(line.strip(), level=1); i += 1; continue

        # Numbered list
        if re.match(r'^\d+\.', line.strip()):
            list_item(line.strip()); i += 1; continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('• '):
            text = line.strip().lstrip('-•').strip()
            list_item(f"• {strip_md(text)}"); i += 1; continue

        # Normal paragraph
        text = line.strip()
        if text and not text.startswith(('#', '|', '!', '>')):
            normal(text)
        i += 1


# ============================================================
# SECTION 1: COVER (no page number)
# ============================================================
print(">> Generating: Cover Page...")
cover_section = doc.sections[0]
setup_cover_section(cover_section)

centered("LAPORAN PROJECT", 16, bold=True, space_after=0)
centered("AUDIT SISTEM INFORMASI", 16, bold=True, space_after=18)

# Logo SD IT Al-huda
if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(18)
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.line_spacing = 1.0
    run_logo = p_logo.add_run()
    run_logo.add_picture(LOGO_PATH, width=Cm(5), height=Cm(5))

centered("EVALUASI TATA KELOLA TEKNOLOGI INFORMASI", 14, bold=True, space_after=0)
centered("MENGGUNAKAN FRAMEWORK COBIT 2019", 14, bold=True, space_after=0)
centered("PADA SD IT AL-HUDA KELAPA GADING", 14, bold=True, space_after=18)
centered("Disusun untuk Memenuhi Tugas Mata Kuliah", 12, space_after=0)
centered("Audit Sistem Informasi (453)", 12, bold=True, space_after=18)
centered("Disusun Oleh:", 12, bold=True, space_after=8)

centered("Fahruroji (11250101)", 12, space_after=0)
centered("Yusuf (11250102)", 12, space_after=18)

centered("Dosen Pengampu:", 12, bold=True, space_after=0)
centered("Juarni Siregar, S.Pd., M.Kom", 12, bold=True, space_after=18)
centered("PROGRAM STUDI SISTEM INFORMASI (S1)", 12, bold=True, space_after=0)
centered("FAKULTAS TEKNOLOGI INFORMASI", 12, bold=True, space_after=0)
centered("UNIVERSITAS NUSA MANDIRI", 12, bold=True, space_after=0)
centered("JAKARTA", 12, bold=True, space_after=0)
centered("2026", 12, bold=True)


# ============================================================
# SECTION 2: FRONT MATTER (romawi, mulai dari ii)
# ============================================================
print(">> Generating: Front Matter...")
front_section = add_new_section()
setup_front_matter_section(front_section, start_page=2)

# ============================================================
# KATA PENGANTAR
# ============================================================
centered("KATA PENGANTAR", 14, bold=True, space_after=24)

for para in [
    "Puji syukur kami panjatkan kehadirat Tuhan Yang Maha Esa, karena atas berkat dan rahmat-Nya kami dapat menyelesaikan laporan project audit sistem informasi ini dengan baik. Laporan ini disusun untuk memenuhi tugas mata kuliah Audit Sistem Informasi (453) pada Program Studi Sistem Informasi (S1), Fakultas Teknologi Informasi, Universitas Nusa Mandiri.",
    'Laporan ini membahas tentang evaluasi tata kelola teknologi informasi pada SD IT Al-huda Kelapa Gading menggunakan framework COBIT 2019. Melalui audit ini, kami melakukan penilaian terhadap tingkat kapabilitas (capability level) tata kelola TI pada lima domain yang relevan dengan kondisi sekolah, yaitu APO07 (Managed Human Resources), BAI09 (Managed Assets), DSS01 (Managed Operations), DSS03 (Managed Problems), dan DSS05 (Managed Security Services).',
    "Kami menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kami mengharapkan kritik dan saran yang membangun dari Ibu dosen pengampu dan rekan-rekan mahasiswa guna penyempurnaan laporan ini di kemudian hari.",
    "Kami mengucapkan terima kasih kepada Ibu Juarni Siregar, S.Pd., M.Kom selaku dosen pengampu, SD IT Al-huda Kelapa Gading yang telah memberikan kesempatan untuk melakukan audit, seluruh responden yang telah berpartisipasi, dan semua pihak yang telah membantu dalam penyelesaian laporan ini.",
    "Semoga laporan ini dapat memberikan manfaat dan kontribusi dalam pengembangan tata kelola teknologi informasi di SD IT Al-huda Kelapa Gading.",
]:
    normal(para)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
r = p.add_run("Jakarta, Mei 2026"); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tim Penulis"); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
doc.add_page_break()


# ============================================================
# DAFTAR ISI
# ============================================================
centered("DAFTAR ISI", 14, bold=True, space_after=24)

toc = [
    ("HALAMAN JUDUL", "i", True),
    ("KATA PENGANTAR", "ii", True),
    ("DAFTAR ISI", "iii", True),
    ("DAFTAR TABEL", "iv", True),
    ("DAFTAR GAMBAR", "v", True),
    ("", "", False),
    ("BAB I PENDAHULUAN", "1", True),
    ("A. Latar Belakang", "1", False),
    ("B. Permasalahan", "2", False),
    ("C. Tujuan", "3", False),
    ("", "", False),
    ("BAB II PEMBAHASAN", "4", True),
    ("A. Framework yang Digunakan", "4", False),
    ("B. Domain dan Sub-Domain yang Digunakan", "5", False),
    ("C. Capability Level dan GAP", "6", False),
    ("", "", False),
    ("BAB III PENUTUP", "10", True),
    ("A. Kesimpulan", "10", False),
    ("B. Saran", "11", False),
    ("", "", False),
    ("DAFTAR PUSTAKA", "12", True),
    ("LAMPIRAN", "13", True),
]
for item, page, is_bold in toc:
    if not item:
        doc.add_paragraph()
        continue
    add_toc_entry(item, page, is_bold=is_bold, indent=not is_bold)

doc.add_page_break()


# ============================================================
# DAFTAR TABEL
# ============================================================
centered("DAFTAR TABEL", 14, bold=True, space_after=24)
tables_list = [
    ("Tabel 2.1", "Capability Level COBIT 2019 (Level 0–5)"),
    ("Tabel 2.2", "Rating Scale Ketercapaian"),
    ("Tabel 2.3", "Domain COBIT 2019 yang Dipilih"),
    ("Tabel 2.4", "RACI Matrix"),
    ("Tabel 2.5", "Hasil Perhitungan Capability Level APO07"),
    ("Tabel 2.6", "Hasil Perhitungan Capability Level BAI09"),
    ("Tabel 2.7", "Hasil Perhitungan Capability Level DSS01"),
    ("Tabel 2.8", "Hasil Perhitungan Capability Level DSS03"),
    ("Tabel 2.9", "Hasil Perhitungan Capability Level DSS05"),
    ("Tabel 2.10", "Ringkasan Capability Level dan GAP Analysis"),
]
for num, title in tables_list:
    add_toc_entry(f"{num} \u2014 {title}", "...")

doc.add_page_break()


# ============================================================
# DAFTAR GAMBAR
# ============================================================
centered("DAFTAR GAMBAR", 14, bold=True, space_after=24)
gambar_list = [
    ("Gambar 2.1", "Alur Metodologi Penelitian"),
    ("Gambar 2.2", "Radar Chart Capability Level 5 Domain COBIT 2019"),
    ("Gambar 2.3", "GAP Analysis — As-is vs To-be per Domain"),
    ("Gambar 2.4", "Persentase Ketercapaian per Level per Domain"),
]
for num, title in gambar_list:
    add_toc_entry(f"{num} \u2014 {title}", "...")




# ============================================================
# SECTION 3: BAB I - PENDAHULUAN (arabic, restart at 1)
# ============================================================
print(">> Generating: BAB I...")
bab1_section = add_new_section()
setup_bab_section(bab1_section, restart_at=1)

centered("BAB I", 14, bold=True, space_after=0)
centered("PENDAHULUAN", 14, bold=True, space_after=24)

bab1_path = os.path.join(LAPORAN_DIR, "BAB-I_Pendahuluan.md")
bab1 = open(bab1_path, 'r', encoding='utf-8').read()
process_md_content(bab1)


# ============================================================
# SECTION 4: BAB II - PEMBAHASAN (arabic, continue)
# ============================================================
print(">> Generating: BAB II...")
bab2_section = add_new_section()
setup_bab_section(bab2_section)

centered("BAB II", 14, bold=True, space_after=0)
centered("PEMBAHASAN", 14, bold=True, space_after=24)

bab2_path = os.path.join(LAPORAN_DIR, "BAB-II_Pembahasan.md")
bab2 = open(bab2_path, 'r', encoding='utf-8').read()
process_md_content(bab2)


# ============================================================
# SECTION 5: BAB III - PENUTUP (arabic, continue)
# ============================================================
print(">> Generating: BAB III...")
bab3_section = add_new_section()
setup_bab_section(bab3_section)

centered("BAB III", 14, bold=True, space_after=0)
centered("PENUTUP", 14, bold=True, space_after=24)

bab3_path = os.path.join(LAPORAN_DIR, "BAB-III_Penutup.md")
bab3 = open(bab3_path, 'r', encoding='utf-8').read()
process_md_content(bab3)


# ============================================================
# SECTION 6: DAFTAR PUSTAKA (arabic, continue, footer center)
# ============================================================
print(">> Generating: Daftar Pustaka...")
dp_section = add_new_section()
setup_daftar_pustaka_section(dp_section)

centered("DAFTAR PUSTAKA", 14, bold=True, space_after=24)

refs_path = os.path.join(LAPORAN_DIR, "Daftar_Pustaka.md")
refs_text = open(refs_path, 'r', encoding='utf-8').read()
for line in refs_text.split('\n'):
    line = line.strip()
    if not line or line.startswith('#'): continue
    # Handle italic book/journal titles
    text = re.sub(r'\*(.+?)\*', r'\1', line)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"[OK] Laporan Word berhasil disimpan: {OUTPUT_PATH}")
print(f"[i] Format: TNR 12pt, Spasi 1.5, Margin 4-3-2.5-2.5 cm")
print(f"[i] Pagination: Cover=hidden, Front matter=romawi, BAB=arabic")
