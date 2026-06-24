import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

BASE_DIR = "/www/wwwroot/audit-si-sdit/TUGAS-AUDIT-SDIT/04_Presentasi"
INPUT_MD = os.path.join(BASE_DIR, "Panduan_Presentasi.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "Panduan_Presentasi.docx")

doc = Document()

# Setup margins
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

# Add Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("PANDUAN PRESENTASI & CATATAN SPEAKER")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Audit Sistem Informasi SD IT Al-huda (COBIT 2019)")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)
run2.font.bold = True

doc.add_paragraph() # spacing

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text

def add_paragraph(text, style='Normal'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    # Handle bold markdown
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.bold = True
        else:
            italic_parts = re.split(r'(\*[^*]+?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    r = p.add_run(ipart[1:-1])
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                    r.font.italic = True
                else:
                    r = p.add_run(ipart)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
    return p

with open(INPUT_MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line or line == '---' or line.startswith('# PANDUAN'):
        continue
    
    if line.startswith('## '):
        doc.add_paragraph() # space before heading
        p = doc.add_paragraph()
        r = p.add_run(line[3:])
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = docx.shared.RGBColor(0, 100, 0) # Dark green
        p.paragraph_format.space_after = Pt(4)
        
    elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.'):
        p = add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1.25)
    elif line.startswith('- '):
        p = add_paragraph("• " + line[2:])
        p.paragraph_format.left_indent = Cm(1.25)
    else:
        add_paragraph(line)

doc.save(OUTPUT_DOCX)
print(f"File panduan berhasil digenerate di: {OUTPUT_DOCX}")
