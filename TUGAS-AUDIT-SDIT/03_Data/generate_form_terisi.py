"""
Generate Form Kuesioner TERISI per Responden — Audit SI SD IT Al-huda Kelapa Gading
Menghasilkan 8 file DOCX (1 per responden) dengan jawaban sudah terisi (✓)
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = "/www/wwwroot/audit-si-sdit/TUGAS-AUDIT-SDIT/03_Data"
OUTPUT_DIR = os.path.join(BASE_DIR, "form_responden")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============================================================
# DATA RESPONDEN & JAWABAN
# ============================================================
RESPONDENTS = [
    {
        'code': 'R1', 'initial': 'FA', 'name': 'Faisal Akbar, M.Pd.',
        'jabatan': 'Kepala Sekolah', 'divisi': 'Manajemen',
        'lama': '8 tahun', 'pendidikan': 'S2 Manajemen Pendidikan',
        'tanggal': '12 / 05 / 2026',
    },
    {
        'code': 'R2', 'initial': 'SR', 'name': 'Siti Rohmah, S.Pd.I.',
        'jabatan': 'Wakil Kepala Sekolah', 'divisi': 'Manajemen',
        'lama': '5 tahun', 'pendidikan': 'S1 Pendidikan Agama Islam',
        'tanggal': '12 / 05 / 2026',
    },
    {
        'code': 'R3', 'initial': 'BD', 'name': 'Budi Darmawan, S.E.',
        'jabatan': 'Kepala Tata Usaha', 'divisi': 'Administrasi',
        'lama': '4 tahun', 'pendidikan': 'S1 Administrasi Perkantoran',
        'tanggal': '13 / 05 / 2026',
    },
    {
        'code': 'R4', 'initial': 'AH', 'name': 'Arif Hidayat, A.Md.',
        'jabatan': 'Operator Sekolah', 'divisi': 'IT & Data',
        'lama': '3 tahun', 'pendidikan': 'D3 Teknik Komputer',
        'tanggal': '13 / 05 / 2026',
    },
    {
        'code': 'R5', 'initial': 'NW', 'name': 'Nurul Wulandari, S.Pd.',
        'jabatan': 'Guru Wali Kelas', 'divisi': 'Akademik',
        'lama': '3 tahun', 'pendidikan': 'S1 PGSD',
        'tanggal': '14 / 05 / 2026',
    },
    {
        'code': 'R6', 'initial': 'RA', 'name': 'Reza Aditya, S.Kom.',
        'jabatan': 'Guru TIK', 'divisi': 'Akademik / IT',
        'lama': '2 tahun', 'pendidikan': 'S1 Pendidikan TIK',
        'tanggal': '14 / 05 / 2026',
    },
    {
        'code': 'R7', 'initial': 'DI', 'name': 'Dina Indriani, S.E.',
        'jabatan': 'Bendahara BOSP', 'divisi': 'Keuangan',
        'lama': '4 tahun', 'pendidikan': 'S1 Akuntansi',
        'tanggal': '15 / 05 / 2026',
    },
    {
        'code': 'R8', 'initial': 'MA', 'name': 'Maulana Asyraf, S.IP.',
        'jabatan': 'Pustakawan', 'divisi': 'Perpustakaan',
        'lama': '3 tahun', 'pendidikan': 'S1 Ilmu Perpustakaan',
        'tanggal': '15 / 05 / 2026',
    },
]

# Jawaban per responden per domain (sesuai Master Excel)
# Format: ANSWERS[domain_idx][question_idx] = [R1, R2, ..., R8]
ANSWERS = {
    'APO07': [
        # Level 1
        [4, 4, 3, 4, 3, 4, 3, 3],  # Q1
        [3, 3, 3, 4, 3, 3, 3, 2],  # Q2
        [3, 3, 3, 3, 3, 3, 3, 2],  # Q3
        # Level 2
        [3, 3, 3, 3, 3, 2, 2, 2],  # Q4
        [3, 3, 3, 3, 2, 2, 2, 2],  # Q5
        [4, 3, 3, 3, 3, 2, 2, 2],  # Q6
        # Level 3
        [3, 3, 2, 2, 2, 2, 2, 2],  # Q7
        [3, 3, 2, 2, 2, 2, 2, 2],  # Q8
        [3, 2, 2, 2, 2, 2, 2, 2],  # Q9
    ],
    'BAI09': [
        [4, 4, 3, 4, 4, 3, 3, 3],
        [4, 3, 4, 4, 3, 3, 3, 3],
        [3, 3, 4, 3, 3, 3, 3, 2],
        [3, 4, 3, 3, 3, 3, 2, 2],
        [3, 3, 4, 3, 3, 2, 2, 2],
        [3, 3, 3, 3, 2, 3, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
    ],
    'DSS01': [
        [4, 4, 4, 4, 3, 4, 3, 3],
        [4, 4, 3, 4, 3, 3, 3, 2],
        [4, 3, 3, 3, 3, 3, 3, 2],
        [3, 3, 4, 3, 3, 3, 3, 2],
        [3, 3, 3, 4, 3, 3, 2, 2],
        [3, 3, 3, 3, 3, 2, 2, 2],
        [3, 3, 3, 2, 2, 2, 2, 2],
        [3, 3, 3, 2, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
    ],
    'DSS03': [
        [4, 4, 4, 4, 3, 3, 3, 3],
        [4, 4, 3, 3, 3, 3, 3, 3],
        [4, 3, 4, 3, 3, 3, 3, 2],
        [3, 3, 3, 4, 3, 3, 3, 2],
        [3, 4, 3, 3, 3, 3, 2, 2],
        [3, 3, 3, 3, 3, 2, 2, 2],
        [3, 3, 3, 2, 2, 2, 2, 2],
        [3, 3, 3, 2, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
    ],
    'DSS05': [
        [4, 4, 4, 4, 4, 3, 3, 3],
        [4, 4, 3, 4, 3, 3, 3, 3],
        [4, 4, 3, 3, 3, 3, 3, 2],
        [3, 3, 4, 3, 3, 2, 2, 2],
        [3, 3, 3, 4, 2, 2, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
        [3, 3, 2, 2, 2, 2, 2, 2],
        [3, 2, 2, 2, 2, 2, 2, 2],
    ],
}

# Domain descriptions + questions (same as form template)
DOMAINS = [
    {
        'code': 'APO07',
        'name': 'Managed Human Resources (Pengelolaan SDM)',
        'desc': 'Domain ini mengevaluasi kompetensi TIK dari guru dan staf serta program pelatihannya di sekolah.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kompetensi dasar guru dan staf dalam menggunakan perangkat TIK sudah diidentifikasi',
                'Penilaian kinerja guru terkait pemanfaatan TIK dalam pembelajaran dilakukan secara rutin',
                'Kebutuhan pelatihan TIK untuk guru dan staf didokumentasikan dengan baik',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan pelatihan TIK yang terstruktur untuk meningkatkan kompetensi guru',
                'Keikutsertaan guru dan staf dalam pelatihan TIK dipantau dan dievaluasi',
                'Manajemen sekolah memberikan dukungan fasilitas untuk pengembangan kompetensi TIK',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar dan indikator kinerja untuk evaluasi kompetensi TIK',
                'Program sertifikasi atau pelatihan TIK diwajibkan bagi guru secara berkala',
                'Evaluasi efektivitas pelatihan TIK dilakukan secara komprehensif',
            ],
        },
    },
    {
        'code': 'BAI09',
        'name': 'Managed Assets (Pengelolaan Aset TIK)',
        'desc': 'Domain ini mengevaluasi pengelolaan dan pemeliharaan perangkat keras (hardware) dan lunak (software) di sekolah, khususnya Lab Komputer.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Seluruh aset perangkat keras (PC Lab, proyektor, dll) dan lunak telah didata',
                'Pengecekan kondisi fisik aset TIK dilakukan secara berkala',
                'Terdapat catatan mengenai kerusakan atau perbaikan aset TIK',
            ],
            'Level 2 (Managed Process)': [
                'Pengadaan aset TIK direncanakan dan disesuaikan dengan kebutuhan sekolah',
                'Penggunaan aset TIK diawasi dan dikelola untuk mencegah penyalahgunaan',
                'Evaluasi kelayakan fungsi aset TIK dilakukan secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP tertulis mengenai pemeliharaan dan penghapusan aset TIK yang rusak',
                'Pengelolaan aset TIK dilakukan menggunakan sistem inventaris digital',
                'Audit terhadap kelengkapan dan kondisi aset TIK dilakukan secara berkala',
            ],
        },
    },
    {
        'code': 'DSS01',
        'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi kegiatan operasional harian fasilitas TIK seperti penjadwalan Lab Komputer dan backup data.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Lab komputer dan fasilitas TIK lainnya dapat digunakan sesuai jadwal yang ditetapkan',
                'Terdapat panduan dasar mengenai penggunaan fasilitas TIK di sekolah',
                'Pencadangan data sekolah (seperti data siswa dan nilai) dilakukan',
            ],
            'Level 2 (Managed Process)': [
                'Operasional Lab Komputer dan sistem TU dikelola dan dipantau secara terstruktur',
                'Kinerja dan ketersediaan fasilitas TIK dievaluasi secara berkala',
                'Prosedur pencadangan data sekolah dilakukan secara rutin dan terjadwal',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP resmi mengenai operasional harian fasilitas TIK di sekolah',
                'Pengelolaan operasional TIK diterapkan secara konsisten oleh seluruh pihak',
                'Uji coba pemulihan data (restore data) dari cadangan dilakukan secara periodik',
            ],
        },
    },
    {
        'code': 'DSS03',
        'name': 'Managed Problems (Pengelolaan Masalah)',
        'desc': 'Domain ini mengevaluasi penanganan kendala teknis TIK (seperti saat ujian sekolah, KBM, atau sistem TU eror).',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kendala teknis saat KBM atau ujian dicatat ketika terjadi',
                'Terdapat penanggung jawab (Operator/Guru TIK) yang menangani kendala teknis',
                'Upaya perbaikan atas kendala teknis dilakukan dengan segera',
            ],
            'Level 2 (Managed Process)': [
                'Masalah TIK diklasifikasikan berdasarkan urgensinya (mis. ujian vs proyektor kelas)',
                'Penyelesaian kendala teknis dipantau dan dilaporkan kepada Kepala Sekolah',
                'Evaluasi terhadap kendala TIK yang sering terjadi dilakukan untuk perbaikan',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP untuk eskalasi dan penyelesaian masalah TIK',
                'Waktu respons dan penyelesaian masalah (SLA) telah ditetapkan dan diukur',
                'Database riwayat penyelesaian masalah didokumentasikan untuk referensi',
            ],
        },
    },
    {
        'code': 'DSS05',
        'name': 'Managed Security Services (Pengelolaan Keamanan Layanan)',
        'desc': 'Domain ini mengevaluasi aspek keamanan data siswa, administrasi keuangan (BOSP), dan jaringan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perangkat komputer sekolah telah dilengkapi dengan perlindungan dasar (antivirus)',
                'Akses ke sistem administrasi sekolah dilindungi oleh kata sandi (password)',
                'Sosialisasi mengenai pentingnya menjaga kerahasiaan akun sudah diberikan',
            ],
            'Level 2 (Managed Process)': [
                'Hak akses ke data sensitif (seperti nilai dan Dapodik) dikelola dan dibatasi',
                'Pemantauan terhadap upaya akses yang tidak sah atau aktivitas mencurigakan dilakukan',
                'Keamanan jaringan sekolah (seperti WiFi) dikelola dan dilindungi',
            ],
            'Level 3 (Established Process)': [
                'Terdapat kebijakan tertulis mengenai keamanan informasi dan privasi data',
                'Audit terhadap pengaturan keamanan dan hak akses dilakukan secara berkala',
                'Pelatihan kesadaran keamanan informasi diberikan secara rutin kepada staf',
            ],
        },
    },
]


# ============================================================
# GENERATE PER RESPONDEN
# ============================================================
def generate_form(resp_idx, resp):
    """Generate satu DOCX form terisi untuk satu responden"""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    def centered(text, size=12, bold=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold

    def normal(text, bold=False, indent=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold

    def set_cell_font(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, bg=None):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
            if color:
                r.font.color.rgb = color
        if bg:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # ── COVER ──
    centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
    centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
    centered("SD IT AL-HUDA KELAPA GADING", 14, bold=True, space_after=6)
    centered("Framework COBIT 2019", 12, space_after=18)
    centered(f"— Form Isian Responden {resp['code']} —", 12, bold=True, space_after=24)

    # ── IDENTITAS RESPONDEN (TERISI) ──
    centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=12)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    fields = [
        ("Nama", resp['name']),
        ("Jabatan", resp['jabatan']),
        ("Unit / Bidang", resp['divisi']),
        ("Lama Mengabdi", resp['lama']),
        ("Pendidikan Terakhir", resp['pendidikan']),
        ("Tanggal Pengisian", resp['tanggal']),
    ]
    for i, (label, value) in enumerate(fields):
        set_cell_font(table.rows[i].cells[0], label, 11)
        set_cell_font(table.rows[i].cells[1], ":", 11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(table.rows[i].cells[2], value, 11, bold=True)
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(0.5)
        table.rows[i].cells[2].width = Cm(9)

    doc.add_page_break()

    # ── KUESIONER PER DOMAIN (JAWABAN TERISI) ──
    q_global = 0
    for d_idx, domain in enumerate(DOMAINS):
        centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
        centered(domain['name'], 12, bold=True, space_after=6)
        normal(domain['desc'], space_after=12)

        for level_name, questions in domain['levels'].items():
            normal(level_name, bold=True, space_after=6)

            num_q = len(questions)
            table = doc.add_table(rows=1 + num_q, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
                set_cell_font(table.rows[0].cells[j], h, 10, bold=True,
                             align=WD_ALIGN_PARAGRAPH.CENTER,
                             color=RGBColor(255, 255, 255), bg="1B3A5C")

            # Questions with filled answers
            for qi, question in enumerate(questions):
                row = table.rows[qi + 1]
                q_num = q_global + 1

                set_cell_font(row.cells[0], str(q_num), 10, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(row.cells[1], question, 10)

                # Get this respondent's answer
                answer = ANSWERS[domain['code']][q_global % 9][resp_idx]

                # Fill checkmark on the correct column
                for ci in range(2, 7):
                    score = ci - 1  # col 2=1(STS), col 3=2(TS), col 4=3(R), col 5=4(S), col 6=5(SS)
                    if score == answer:
                        set_cell_font(row.cells[ci], "✓", 12, bold=True,
                                     align=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        row.cells[ci].text = ""
                        row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[ci].width = Cm(1.5)

                row.cells[0].width = Cm(1)
                row.cells[1].width = Cm(8)
                q_global += 1

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        if d_idx < len(DOMAINS) - 1:
            doc.add_page_break()

    # ── FOOTER ──
    doc.add_paragraph()
    normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"( {resp['name']} )"); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True

    # Save
    filename = f"Kuesioner_Terisi_{resp['code']}_{resp['initial']}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filename


# ============================================================
# MAIN
# ============================================================
print("=" * 60)
print("  GENERATE FORM KUESIONER TERISI PER RESPONDEN (SD IT Al-huda)")
print("=" * 60)

files = []
for idx, resp in enumerate(RESPONDENTS):
    fname = generate_form(idx, resp)
    files.append(fname)
    print(f"  ✅ {resp['code']} ({resp['initial']}) — {resp['jabatan']:30s} → {fname}")

print(f"\n📂 Output folder: {OUTPUT_DIR}")
print(f"📄 Total file: {len(files)} form kuesioner terisi")
print(f"📝 Setiap form berisi: Identitas (terisi) + 45 soal (jawaban ✓)")
