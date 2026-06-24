"""
Generate Form Kuesioner Word (.docx) — Audit SI SD IT Al-huda Kelapa Gading
Form ini bisa dicetak atau dikirim ke responden untuk diisi
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = "/www/wwwroot/audit-si-sdit/TUGAS-AUDIT-SDIT/03_Data"
OUTPUT = os.path.join(BASE_DIR, "Form_Kuesioner_Audit_SI_SDIT.docx")

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def normal(text, bold=False, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold
    return p


# ============================================================
# HALAMAN 1: COVER & IDENTITAS
# ============================================================
centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
centered("SD IT AL-HUDA KELAPA GADING", 14, bold=True, space_after=18)
centered("Framework COBIT 2019", 12, space_after=24)

normal("Kepada Yth. Bapak/Ibu Responden", bold=True, space_after=6)
normal("di Tempat", space_after=12)

normal("Dengan hormat,", indent=True, space_after=6)
normal("Perkenalkan, kami mahasiswa Program Studi Sistem Informasi (S1), Fakultas Teknologi Informasi, Universitas Nusa Mandiri yang sedang melakukan penelitian dalam rangka memenuhi tugas mata kuliah Audit Sistem Informasi. Penelitian ini bertujuan untuk mengevaluasi tata kelola teknologi informasi pada SD IT Al-huda Kelapa Gading menggunakan framework COBIT 2019.", indent=True, space_after=6)
normal("Kami mengharapkan kesediaan Bapak/Ibu untuk mengisi kuesioner ini dengan sebaik-baiknya. Data yang diperoleh akan digunakan semata-mata untuk kepentingan akademis dan dijamin kerahasiaannya.", indent=True, space_after=6)
normal("Atas perhatian dan kesediaan Bapak/Ibu, kami ucapkan terima kasih.", indent=True, space_after=12)

# Tim peneliti
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(12)
r = p.add_run("Hormat kami,"); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tim Peneliti"); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

doc.add_page_break()


# ============================================================
# HALAMAN 2: IDENTITAS RESPONDEN
# ============================================================
centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=18)

normal("Petunjuk: Mohon lengkapi data berikut.", space_after=12)

# Identity table
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
fields = [
    ("Nama", ""),
    ("Jabatan", ""),
    ("Unit / Bidang", ""),
    ("Lama Mengabdi", "...... tahun"),
    ("Pendidikan Terakhir", ""),
    ("Tanggal Pengisian", "...... / ...... / 2026"),
]
for i, (label, default) in enumerate(fields):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = ":"
    table.rows[i].cells[2].text = default
    for j in range(3):
        p = table.rows[i].cells[j].paragraphs[0]
        p.runs[0].font.name = 'Times New Roman' if p.runs else None
        p.runs[0].font.size = Pt(12) if p.runs else None
    table.rows[i].cells[0].width = Cm(5)
    table.rows[i].cells[1].width = Cm(0.5)
    table.rows[i].cells[2].width = Cm(9)

doc.add_paragraph().paragraph_format.space_after = Pt(18)

# Petunjuk pengisian
centered("PETUNJUK PENGISIAN", 14, bold=True, space_after=12)
normal("Berilah tanda centang (✓) pada salah satu kolom yang sesuai dengan pendapat Bapak/Ibu terhadap pernyataan-pernyataan berikut:", space_after=12)

# Scale explanation
scale_table = doc.add_table(rows=6, cols=3)
scale_table.style = 'Table Grid'
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for j, h in enumerate(["Skor", "Singkatan", "Keterangan"]):
    cell = scale_table.rows[0].cells[j]
    cell.text = h
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    for r in p.runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

scales = [
    ("1", "STS", "Sangat Tidak Setuju"),
    ("2", "TS", "Tidak Setuju"),
    ("3", "R", "Ragu-ragu / Netral"),
    ("4", "S", "Setuju"),
    ("5", "SS", "Sangat Setuju"),
]
for i, (skor, singkat, ket) in enumerate(scales):
    for j, val in enumerate([skor, singkat, ket]):
        cell = scale_table.rows[i+1].cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)

doc.add_page_break()


# ============================================================
# HALAMAN 3+: KUESIONER PER DOMAIN
# ============================================================
domains = [
    {
        'code': 'APO07',
        'name': 'Managed Human Resources (Pengelolaan SDM)',
        'desc': 'Domain ini mengevaluasi kompetensi TIK dari guru dan staf serta program pelatihannya di sekolah.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kompetensi dasar guru dan staf dalam menggunakan perangkat TIK sudah diidentifikasi',
                'Penilaian kinerja guru terkait pemanfaatan TIK dalam pembelajaran dilakukan secara rutin',
                'Kebutuhan pelatihan TIK untuk guru dan staf didokumentasikan dengan baik',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan pelatihan TIK yang terstruktur untuk meningkatkan kompetensi guru',
                'Keikutsertaan guru dan staf dalam pelatihan TIK dipantau dan dievaluasi',
                'Manajemen sekolah memberikan dukungan fasilitas untuk pengembangan kompetensi TIK',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar dan indikator kinerja untuk evaluasi kompetensi TIK',
                'Program sertifikasi atau pelatihan TIK diwajibkan bagi guru secara berkala',
                'Evaluasi efektivitas pelatihan TIK dilakukan secara komprehensif',
            ],
        },
    },
    {
        'code': 'BAI09',
        'name': 'Managed Assets (Pengelolaan Aset TIK)',
        'desc': 'Domain ini mengevaluasi pengelolaan dan pemeliharaan perangkat keras (hardware) dan lunak (software) di sekolah, khususnya Lab Komputer.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Seluruh aset perangkat keras (PC Lab, proyektor, dll) dan lunak telah didata',
                'Pengecekan kondisi fisik aset TIK dilakukan secara berkala',
                'Terdapat catatan mengenai kerusakan atau perbaikan aset TIK',
            ],
            'Level 2 (Managed Process)': [
                'Pengadaan aset TIK direncanakan dan disesuaikan dengan kebutuhan sekolah',
                'Penggunaan aset TIK diawasi dan dikelola untuk mencegah penyalahgunaan',
                'Evaluasi kelayakan fungsi aset TIK dilakukan secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP tertulis mengenai pemeliharaan dan penghapusan aset TIK yang rusak',
                'Pengelolaan aset TIK dilakukan menggunakan sistem inventaris digital',
                'Audit terhadap kelengkapan dan kondisi aset TIK dilakukan secara berkala',
            ],
        },
    },
    {
        'code': 'DSS01',
        'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi kegiatan operasional harian fasilitas TIK seperti penjadwalan Lab Komputer dan backup data.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Lab komputer dan fasilitas TIK lainnya dapat digunakan sesuai jadwal yang ditetapkan',
                'Terdapat panduan dasar mengenai penggunaan fasilitas TIK di sekolah',
                'Pencadangan data sekolah (seperti data siswa dan nilai) dilakukan',
            ],
            'Level 2 (Managed Process)': [
                'Operasional Lab Komputer dan sistem TU dikelola dan dipantau secara terstruktur',
                'Kinerja dan ketersediaan fasilitas TIK dievaluasi secara berkala',
                'Prosedur pencadangan data sekolah dilakukan secara rutin dan terjadwal',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP resmi mengenai operasional harian fasilitas TIK di sekolah',
                'Pengelolaan operasional TIK diterapkan secara konsisten oleh seluruh pihak',
                'Uji coba pemulihan data (restore data) dari cadangan dilakukan secara periodik',
            ],
        },
    },
    {
        'code': 'DSS03',
        'name': 'Managed Problems (Pengelolaan Masalah)',
        'desc': 'Domain ini mengevaluasi penanganan kendala teknis TIK (seperti saat ujian sekolah, KBM, atau sistem TU eror).',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kendala teknis saat KBM atau ujian dicatat ketika terjadi',
                'Terdapat penanggung jawab (Operator/Guru TIK) yang menangani kendala teknis',
                'Upaya perbaikan atas kendala teknis dilakukan dengan segera',
            ],
            'Level 2 (Managed Process)': [
                'Masalah TIK diklasifikasikan berdasarkan urgensinya (mis. ujian vs proyektor kelas)',
                'Penyelesaian kendala teknis dipantau dan dilaporkan kepada Kepala Sekolah',
                'Evaluasi terhadap kendala TIK yang sering terjadi dilakukan untuk perbaikan',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP untuk eskalasi dan penyelesaian masalah TIK',
                'Waktu respons dan penyelesaian masalah (SLA) telah ditetapkan dan diukur',
                'Database riwayat penyelesaian masalah didokumentasikan untuk referensi',
            ],
        },
    },
    {
        'code': 'DSS05',
        'name': 'Managed Security Services (Pengelolaan Keamanan Layanan)',
        'desc': 'Domain ini mengevaluasi aspek keamanan data siswa, administrasi keuangan (BOSP), dan jaringan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perangkat komputer sekolah telah dilengkapi dengan perlindungan dasar (antivirus)',
                'Akses ke sistem administrasi sekolah dilindungi oleh kata sandi (password)',
                'Sosialisasi mengenai pentingnya menjaga kerahasiaan akun sudah diberikan',
            ],
            'Level 2 (Managed Process)': [
                'Hak akses ke data sensitif (seperti nilai dan Dapodik) dikelola dan dibatasi',
                'Pemantauan terhadap upaya akses yang tidak sah atau aktivitas mencurigakan dilakukan',
                'Keamanan jaringan sekolah (seperti WiFi) dikelola dan dilindungi',
            ],
            'Level 3 (Established Process)': [
                'Terdapat kebijakan tertulis mengenai keamanan informasi dan privasi data',
                'Audit terhadap pengaturan keamanan dan hak akses dilakukan secara berkala',
                'Pelatihan kesadaran keamanan informasi diberikan secara rutin kepada staf',
            ],
        },
    },
]

q_global = 1
for d_idx, domain in enumerate(domains):
    centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
    centered(domain['name'], 12, bold=True, space_after=6)
    normal(domain['desc'], space_after=12)

    for level_name, questions in domain['levels'].items():
        normal(level_name, bold=True, space_after=6)

        # Create table for questions
        num_q = len(questions)
        table = doc.add_table(rows=1 + num_q, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
            cell = table.rows[0].cells[j]
            cell.text = h
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Questions
        for qi, question in enumerate(questions):
            row = table.rows[qi + 1]
            # No
            row.cells[0].text = str(q_global)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in row.cells[0].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)

            # Question text
            row.cells[1].text = question
            for r in row.cells[1].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)

            # Empty checkboxes (STS-SS)
            for ci in range(2, 7):
                row.cells[ci].text = ""
                row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set fixed width for checkbox columns
                row.cells[ci].width = Cm(1.5)

            # Set column widths
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(8)

            q_global += 1

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    if d_idx < len(domains) - 1:
        doc.add_page_break()

# Footer note
doc.add_paragraph()
normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True, space_after=6)
normal("Data yang Bapak/Ibu berikan akan dijaga kerahasiaannya dan hanya digunakan untuk kepentingan akademis.", space_after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("(__________________________)"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"✅ Form Kuesioner berhasil disimpan: {OUTPUT}")
print(f"📋 Total domain: 5 (APO07, BAI09, DSS01, DSS03, DSS05)")
print(f"📝 Total pertanyaan: {q_global - 1} soal (9 per domain × 5 domain = 45)")
print(f"📄 Isi: Cover + Identitas Responden + Petunjuk + Kuesioner 5 Domain")
